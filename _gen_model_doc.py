from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

out = Path(r'C:\Users\Administrator\Documents\xwechat_files\wxid_7b0gbpwkxyeq22_4bc7\msg\file\2026-07\模型选型_准确率优先版.docx')

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)

sec = doc.sections[0]
sec.top_margin = Pt(72)
sec.bottom_margin = Pt(72)
sec.left_margin = Pt(90)
sec.right_margin = Pt(90)

def set_run(run, bold=False, size=12):
    run.bold = bold
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)

def add_para(text='', bold=False, size=12, first_line=True):
    p = doc.add_paragraph()
    if first_line:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    set_run(r, bold, size)
    return p

def add_multirun(parts, first_line=True):
    p = doc.add_paragraph()
    if first_line:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    for text, bold in parts:
        r = p.add_run(text)
        set_run(r, bold, 12)
    return p

t = doc.add_paragraph()
t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = t.add_run('手语识别模型选型（准确率优先版）')
set_run(r, True, 16)

add_para('结合本项目“基于手语数字人的无障碍智能交互系统”的建设目标，模型选型应优先考虑公开基准上的识别准确率或翻译效果，同时兼顾代码可获取性、已有权重支持情况、系统集成可行性以及后续优化空间。与单纯追求“是否容易复现”不同，本次筛选将“复现后是否能够达到较高性能”作为首要标准，因此优先保留在公开数据集上结果领先、且已有官方代码或权重支持的模型。')

models = [
    (
        'CVPR 2024“SignGraph”优先推荐',
        'https://github.com/gswycf/SignGraph',
        'SignGraph: A Sign Sequence is Worth Graphs of Nodes',
        'Shiwei Gan, Yafeng Yin, Zhiwei Jiang, Hongkai Wen, Lei Xie, Sanglu Lu',
        'CVPR 2024',
        [
            '从准确率看，SignGraph 是这批候选里非常值得重点考虑的模型。官方仓库给出的结果显示，其在 Phoenix-2014 数据集上测试集 WER 可达到 18.17，在 Phoenix-2014T 上测试集 WER 为 19.44，在不额外引入多模态辅助信息的情况下，已经达到非常强的识别性能。对于连续手语识别任务而言，WER 越低越好，因此该模型在准确率维度上具有明显优势。',
            '从方法上看，SignGraph 将手语序列表示成图结构，通过局部图和时间图建模帧内区域关系与跨帧动态关系，相比传统 CNN 网格建模方式，更贴合手语动作识别的本质需求。这意味着它不仅结果强，而且模型设计本身也有较好的论文阐述价值。',
            '从项目适配性看，SignGraph 的一个重要优点是“高准确率且不依赖额外 cue”，这使得后续接入你们的视频识别链路时更直接，不必强绑定复杂的姿态流或额外传感器。若项目把识别准确率作为最重要指标，SignGraph 非常适合作为首选主模型。',
        ],
    ),
    (
        'NeurIPS 2022“TwoStream Network”优先推荐',
        'https://github.com/FangyunWei/SLRT',
        'Two-Stream Network for Sign Language Recognition and Translation',
        'Yutong Chen, Ronglai Zuo, Fangyun Wei, Yu Wu, Shujie Liu, Brian Mak',
        'NeurIPS 2022',
        [
            'TwoStream Network 是当前手语识别与翻译中兼顾性能和成熟度的代表性模型。官方结果表明，其在 Phoenix-2014 上测试集 WER 为 18.8，在 Phoenix-2014T 上测试集 WER 为 19.3，同时在翻译任务上还可在 Phoenix-2014T 上取得约 29.0 的 BLEU-4，整体表现十分突出。',
            '该模型采用 RGB 视频流与人体关键点流双分支联合建模，能够同时捕获外观信息与结构动作信息，因此在准确率上通常优于单纯视频模型。对于你们项目而言，这种双流设计也与“视频识别 + 骨骼特征 + 数字人驱动”的整体路线较一致。',
            '若单从性能稳定性和公开资料完整度来看，TwoStream 是非常稳妥的高精度方案。它的缺点是推理链路相对更复杂，但如果项目目标强调“结果尽量强”，那么 TwoStream 仍应排在最优先队列。',
        ],
    ),
    (
        'CVPR 2023“CVT-SLR”高精度候选',
        'https://github.com/binbinjiang/CVT-SLR',
        'CVT-SLR: Contrastive Visual-Textual Transformation for Sign Language Recognition with Variational Alignment',
        'Jiangbin Zheng, Yile Wang, Cheng Tan, Siyuan Li, Ge Wang, Jun Xia, Yidong Chen, Stan Z. Li',
        'CVPR 2023',
        [
            'CVT-SLR 是一类高性能单流识别模型，重点利用视觉与文本之间的跨模态对齐来提升识别效果。官方仓库显示，其在 Phoenix-2014 上测试集 WER 可达到 20.06 左右，开发集约为 19.80，在不依赖多流额外输入的条件下，已经优于不少较早的多模态方法。',
            '这类方法的突出优势在于：虽然准确率很强，但整体结构比部分复杂多流模型更“干净”，更适合作为高精度单流代表模型。若你们希望在论文或材料里体现“我们不仅考虑双流模型，也对高性能单流方法进行了比较”，CVT-SLR 很适合被列入重点候选。',
            '需要注意的是，CVT-SLR 在工程实现上对数据预处理和训练配置较敏感，复现门槛会略高于传统 baseline。但如果把准确率放在第一位，它仍然比很多“好复现但性能一般”的模型更值得进入最终候选名单。',
        ],
    ),
    (
        'CVPR 2022“C2SLR”高性能稳妥候选',
        'https://github.com/2000ZRL/LCSA_C2SLR_SRM',
        'C2SLR: Consistency-enhanced Continuous Sign Language Recognition',
        'Ronglai Zuo 等',
        'CVPR 2022',
        [
            'C2SLR 是连续手语识别中一条非常稳健的高性能路线。公开结果显示，其在 Phoenix-2014 上测试集 WER 为 20.4，在 Phoenix-2014T 上测试集 WER 也达到 20.4，属于长期被反复引用和比较的强基线。',
            '与一些更新更复杂的方法相比，C2SLR 的绝对指标不是最顶尖，但它的优点在于结果可靠、代码和配置较完整、方法逻辑清晰，在高精度候选中属于“风险较低”的方案。它使用一致性约束和关键点引导注意力来增强表征能力，因此兼顾了性能和可解释性。',
            '如果你们最终希望形成“一个最强模型 + 一个稳健高性能备选 + 一个经典基线”的组合，那么 C2SLR 非常适合放在备选位。它特别适合在复现和对比实验中承担参照模型角色。',
        ],
    ),
    (
        'ICLR 2025“Uni-Sign”前沿候选',
        'https://github.com/ZechengLi19/Uni-Sign',
        'Uni-Sign: Toward Unified Sign Language Understanding at Scale',
        'Zecheng Li, Wengang Zhou, Weichao Zhao, Kepeng Wu, Hezhen Hu, Houqiang Li',
        'ICLR 2025',
        [
            'Uni-Sign 是近年的统一手语理解框架，研究先进性很强，尤其在中文手语相关任务上具有很高吸引力。论文报告指出，该模型在 CSL-Daily 等任务上取得了显著领先的结果，在 gloss-free 设定下 BLEU-4 相比既有方法有明显提升。对于你们项目这种中文手语数字人系统，Uni-Sign 的研究契合度是很高的。',
            '不过需要说明的是，Uni-Sign 的强项更突出体现在统一预训练和翻译能力上，而不是像 Phoenix 系列 CSLR 模型那样可以直接用一个单一 WER 指标横向比较。因此如果评价标准是“公开视频识别准确率/WER 是否最低”，它未必像 SignGraph、TwoStream、CVT-SLR 那样直观。',
            '尽管如此，若项目希望强调“中文场景适配能力”和“前沿先进性”，Uni-Sign 仍然值得保留在最终候选中。更合理的定位不是最稳的第一复现对象，而是高潜力的前沿方案或后续升级方向。',
        ],
    ),
]

for title, link, paper, authors, venue, reasons in models:
    add_para(title, bold=True, size=13, first_line=False)
    add_multirun([('链接：', True), (link, False)], first_line=False)
    add_multirun([('论文：', True), (paper, False)], first_line=False)
    add_multirun([('作者：', True), (authors, False)], first_line=False)
    add_multirun([('会议：', True), (venue, False)], first_line=False)
    add_multirun([('推荐理由：', True)], first_line=False)
    for txt in reasons:
        add_para(txt)

add_para('综合以上候选模型，可以看出若将“复现后准确率”作为第一优先级，则不应仅仅围绕“是否容易跑通”进行选择，而应重点考虑模型在公开基准上的最优结果、方法成熟度以及是否已有权重支持。为便于项目后续决策，现将主要候选模型对比如下。')

table = doc.add_table(rows=1, cols=6)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
headers = ['模型', '会议', '公开结果（代表指标）', '代码/权重', '适合定位', '推荐等级']
for i, txt in enumerate(headers):
    p = table.rows[0].cells[i].paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(txt)
    set_run(r, True, 11)

rows = [
    ('SignGraph', 'CVPR 2024', 'Phoenix-2014 Test WER 18.17', '代码+权重公开', '高精度主模型', '★★★★★'),
    ('TwoStream Network', 'NeurIPS 2022', 'Phoenix-2014 Test WER 18.8；Phoenix-2014T Test WER 19.3', '代码+权重公开', '高精度主模型', '★★★★★'),
    ('CVT-SLR', 'CVPR 2023', 'Phoenix-2014 Test WER 20.06', '代码+权重公开', '高精度单流候选', '★★★★☆'),
    ('C2SLR', 'CVPR 2022', 'Phoenix-2014 Test WER 20.4；Phoenix-2014T Test WER 20.4', '代码+配置公开', '稳健高性能备选', '★★★★☆'),
    ('Uni-Sign', 'ICLR 2025', 'CSL-Daily gloss-free 结果领先，BLEU-4 显著提升', '代码公开', '中文场景前沿候选', '★★★★☆'),
]
for row in rows:
    cells = table.add_row().cells
    for i, txt in enumerate(row):
        p = cells[i].paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = p.add_run(txt)
        set_run(r, False, 10.5)

add_para('最终建议：若你们当前目标是尽可能做出“准确率高、答辩时有说服力”的识别模块，建议优先选择 SignGraph 或 TwoStream Network 作为主模型；若需要再保留一个高性能备选方案，可加入 CVT-SLR 或 C2SLR；Uni-Sign 更适合作为中文手语方向的前沿增强路线或后续升级方案。也就是说，在“准确率优先”的前提下，本项目最推荐的组合是“SignGraph / TwoStream 作为主线，C2SLR 或 CVT-SLR 作为对比，Uni-Sign 作为前沿补充”。')

doc.save(out)
print(out)
